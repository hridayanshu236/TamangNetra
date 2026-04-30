from docx import Document
from backend.services.sentence_splitter import split as split_sentences

async def process_docx(input_path, output_path, src, tgt, translator, glossary, progress_cb):
    doc = Document(input_path)

    # Count all translatable runs for progress
    all_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                all_runs.append(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            all_runs.append(run)

    total = len(all_runs)
    for i, run in enumerate(all_runs):
        sentences = split_sentences(run.text, src)
        translated_parts = []
        for s in sentences:
            s_mod, placeholders = glossary.apply(s)
            t = await translator.translate(s_mod, src, tgt)
            t = glossary.restore(t, placeholders)
            translated_parts.append(t)
            # Learn terms
            for term in glossary.extract_terms(s):
                if term not in glossary.cache:
                    glossary.learn(term, await translator.translate(term, src, tgt))
        run.text = " ".join(translated_parts)
        # run.bold, run.italic, run.font.size are preserved — python-docx keeps them
        await progress_cb(i + 1, total)

    doc.save(output_path)