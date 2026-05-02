import io
import hashlib
import tempfile
import pandas as pd
from typing import List, Dict, Any, Optional, Tuple
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
from fastapi import UploadFile, Depends
import re
import requests as req
import logging
from bs4 import BeautifulSoup

from app.services.translation import get_translation_service, TranslationService
from app.core.config import Settings, get_settings
from app.models.schemas import TranslationRequest

# New PDF Pipeline
from .pdf.extractor import PdfExtractor, PageData
from .pdf.reconstructor import PdfReconstructor
from .pdf.high_fidelity.extractor import HighFidelityPdfExtractor
from .pdf.high_fidelity.translator import HighFidelityPdfTranslator
from .pdf.high_fidelity.reconstructor import HighFidelityPdfReconstructor

try:
    from .pdf.ocr import OcrHandler
except ImportError:
    OcrHandler = None

logger = logging.getLogger(__name__)

# Module-level font cache
_font_cache: dict[str, bytes] = {}

def _get_font_bytes(url: str) -> bytes:
    """Download and cache a font file by URL."""
    if url not in _font_cache:
        logger.info(f"Downloading font from {url}...")
        r = req.get(url, timeout=20)
        r.raise_for_status()
        _font_cache[url] = r.content
        logger.info("Font downloaded and cached.")
    return _font_cache[url]


class DocumentProcessor:
    def __init__(self, translation_service: TranslationService):
        self.translation_service = translation_service
        self.pdf_extractor = PdfExtractor()
        self.pdf_reconstructor = PdfReconstructor(
            font_regular_path="assets/fonts/NotoSansDevanagari-Regular.ttf",
            font_bold_path="assets/fonts/NotoSansDevanagari-Bold.ttf"
        )
        # High Fidelity Pipeline
        self.hf_extractor = HighFidelityPdfExtractor()
        self.hf_translator = HighFidelityPdfTranslator(translation_service)
        self.hf_reconstructor = HighFidelityPdfReconstructor(
            font_regular_path="assets/fonts/NotoSansDevanagari-Regular.ttf"
        )
        self.ocr_handler = OcrHandler() if OcrHandler else None
        # Cache for translation mappings: {file_hash: {src_text: tgt_text}}
        self._translation_cache: Dict[str, Dict[str, str]] = {}

    async def process_pdf(self, file_content: bytes, src_lang: str, tgt_lang: str, progress_callback=None) -> Dict[str, Any]:
        """
        High-Fidelity PDF Processing:
        1. Extract to HTML (Rich structure)
        2. Translate HTML (Preserving tags/math)
        3. Extract segments for UI preview
        """
        try:
            # 1. Convert to rich HTML
            html_content = self.hf_extractor.pdf_to_html(file_content)
            
            # 2. Translate HTML (this populates the Knowledge Graph cache)
            translated_html = await self.hf_translator.translate_html(
                html_content, src_lang, tgt_lang, progress_callback=progress_callback, translate_all=True
            )
            
            # 3. Extract unique segments for the frontend preview
            # We use BeautifulSoup to get the clean text segments AND images
            soup_orig = BeautifulSoup(html_content, "html.parser")
            soup_trans = BeautifulSoup(translated_html, "html.parser")
            
            # Find all content elements (text or img)
            def _get_segments(soup):
                segs = []
                # We iterate through body descendants to find text and images
                for element in soup.body.find_all(['span', 'td', 'img'], recursive=True):
                    if element.name == 'img':
                        # For images, we provide the tag as is (base64)
                        segs.append(str(element))
                    elif element.string and element.string.strip():
                        segs.append(element.string.strip())
                return segs

            orig_texts = _get_segments(soup_orig)
            trans_texts = _get_segments(soup_trans)
            
            # Ensure we have a 1:1 mapping for the preview segments
            min_len = min(len(orig_texts), len(trans_texts))
            results = [{"original": orig_texts[i], "translated": trans_texts[i]} for i in range(min_len)]
            
            return {
                "original": "\n".join(orig_texts),
                "translated": "\n".join(trans_texts),
                "segments": results
            }
        except Exception as e:
            logger.error(f"High-fidelity PDF processing failed, falling back to basic: {e}")
            # FALLBACK to original basic extractor
            return await self._process_pdf_basic(file_content, src_lang, tgt_lang, progress_callback)

    async def _process_pdf_basic(self, file_content: bytes, src_lang: str, tgt_lang: str, progress_callback=None) -> Dict[str, Any]:
        """Original absolute-positioning extraction for PDF."""
        pages_data = self.pdf_extractor.extract(file_content)
        all_text = []
        for page in pages_data:
            for span in page.spans:
                all_text.append(span.text)
            for table in page.tables:
                for row in table.rows:
                    if row:
                        all_text.append(" | ".join([str(c) if c else "" for c in row]))
        
        unique_texts = list(set([t for t in all_text if t.strip()]))
        translated_list = await self.translation_service.batch_translate(
            unique_texts, src_lang, tgt_lang, progress_callback=progress_callback
        )
        return {
            "original": "\n".join(unique_texts),
            "translated": "\n".join(translated_list),
            "segments": [{"original": o, "translated": t} for o, t in zip(unique_texts, translated_list)]
        }

    async def process_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX run-by-run to match reconstruction exactly."""
        doc = DocxDocument(io.BytesIO(file_content))
        run_texts = []
        
        def _collect_runs(paragraphs):
            for para in paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        run_texts.append(run.text)

        for para in doc.paragraphs:
            _collect_runs([para])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _collect_runs(cell.paragraphs)
                    
        return "---EXACT-BLOCK---".join(run_texts)

    async def process_spreadsheet(self, file_content: bytes, file_type: str) -> str:
        """Extract text from CSV or Excel cell-by-cell to match reconstruction."""
        if file_type == 'csv':
            df = pd.read_csv(io.BytesIO(file_content))
        else:
            df = pd.read_excel(io.BytesIO(file_content))
            
        cells = list(df.columns)
        for _, row in df.iterrows():
            for val in row.values:
                cells.append(str(val) if pd.notna(val) else "")
                
        # Filter empty and join with PAGE-BREAK so they are processed as isolated blocks
        cells = [c for c in cells if str(c).strip()]
        return "---PAGE-BREAK---".join(cells)

    async def process_image(self, file_content: bytes) -> str:
        """Apply OCR directly to an image file."""
        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            raise ValueError(f"Failed to process image: {str(e)}")

    async def process_file(self, file_content: bytes, file_name: str, file_size: int, src_lang: str, tgt_lang: str, progress_callback = None) -> Dict[str, Any]:
        """Main dispatcher for processing files. Now supports yielding segments back to the caller."""
        content = file_content
        filename = file_name.lower()
        
        try:
            if filename.endswith(".pdf"):
                pdf_res = await self.process_pdf(content, src_lang, tgt_lang, progress_callback=progress_callback)
                return {
                    "original": pdf_res["original"],
                    "translated": pdf_res["translated"],
                    "segments": pdf_res["segments"],
                    "fileInfo": {"name": file_name, "type": "pdf", "size": file_size}
                }
            else:
                if filename.endswith(".docx"):
                    extracted_text = await self.process_docx(content)
                elif filename.endswith(".csv"):
                    extracted_text = await self.process_spreadsheet(content, 'csv')
                elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                    extracted_text = await self.process_spreadsheet(content, 'excel')
                elif filename.endswith((".jpg", ".jpeg", ".png")):
                    extracted_text = await self.process_image(content)
                else:
                    raise ValueError("Unsupported file format")
                    
                if not extracted_text.strip():
                    raise ValueError("No text could be extracted from the document")
                    
                segments = self._split_into_sentences(extracted_text)
                
                # Translate segments using the enhanced streaming callback
                translated_texts = await self.translation_service.batch_translate(
                    texts=segments,
                    source_language=src_lang,
                    target_language=tgt_lang,
                    progress_callback=progress_callback
                )
                
                results = []
                for orig, trans in zip(segments, translated_texts):
                    results.append({"original": orig, "translated": trans})
                    
                return {
                    "original": extracted_text,
                    "translated": "\n".join(translated_texts),
                    "segments": results,
                    "fileInfo": {
                        "name": file_name,
                        "type": filename.split('.')[-1] if '.' in filename else 'unknown',
                        "size": file_size
                    }
                }
            
        except Exception as e:
            raise ValueError(f"Error processing document: {str(e)}")

    def _split_into_sentences(self, text: str) -> list[str]:
        """
        Split text into sentences.
        We split by blocks (pages) first to ensure 100% cache hits 
        during reconstruction while remaining immune to line breaks.
        """
        if "---EXACT-BLOCK---" in text:
            # DOCX uses exact blocks run-by-run without sentence splitting
            return [b for b in text.split("---EXACT-BLOCK---") if b.strip()]

        all_segments = []
        # Split by our internal page marker if present
        blocks = text.split("---PAGE-BREAK---")
        
        for block in blocks:
            # 1. Normalize ALL whitespace within this block into single spaces
            block = re.sub(r'\s+', ' ', block).strip()
            if not block:
                continue
                
            # 2. Split into sentences using punctuation boundaries
            # Supports English (.!?) and Devanagari (।)
            line_sentences = [s.strip() for s in re.split(r'(?<=[.!?।])\s+', block) if s.strip()]
            
            for s in line_sentences:
                # 3. Standardize sentence endings
                if not re.search(r'[.!?।]$', s):
                    s += '.'
                all_segments.append(s)
                
        return all_segments

    async def reconstruct_file(self, file: UploadFile, src_lang: str, tgt_lang: str) -> tuple[bytes, str]:
        """Reconstruct the translated file with exact layout/formatting."""
        content = await file.read()
        filename = file.filename.lower()
        
        try:
            if filename.endswith(".pdf"):
                return await self._reconstruct_pdf(content, src_lang, tgt_lang), "application/pdf"
            elif filename.endswith(".docx"):
                return await self._reconstruct_docx(content, src_lang, tgt_lang), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif filename.endswith(".csv"):
                return await self._reconstruct_spreadsheet(content, src_lang, tgt_lang, "csv"), "text/csv"
            elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                return await self._reconstruct_spreadsheet(content, src_lang, tgt_lang, "excel"), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            else:
                raise ValueError("Unsupported file format for reconstruction")
        except Exception as e:
            raise ValueError(f"Error reconstructing document: {str(e)}")

    async def _reconstruct_pdf(self, content: bytes, src_lang: str, tgt_lang: str) -> bytes:
        """High-fidelity PDF reconstruction using WeasyPrint (with fallback to ReportLab)."""
        try:
            # 1. Convert to rich HTML
            html_content = self.hf_extractor.pdf_to_html(content)
            
            # 2. Translate HTML (hits the persistent disk cache)
            translated_html = await self.hf_translator.translate_html(
                html_content, src_lang, tgt_lang, progress_callback=None, cache_only=True
            )
            
            # 3. Reconstruct via WeasyPrint
            return self.hf_reconstructor.html_to_pdf(translated_html)
            
        except Exception as e:
            logger.error(f"High-fidelity reconstruction failed, falling back: {e}")
            
            # FALLBACK to original ReportLab-based reconstruction
            pages_data = self.pdf_extractor.extract(content)
            texts_to_translate = []
            for page in pages_data:
                for span in page.spans:
                    texts_to_translate.append(span.text)
                for table in page.tables:
                    for row in table.rows:
                        if row:
                            texts_to_translate.extend([str(c) if c else "" for c in row])

            unique_texts = list(set([t for t in texts_to_translate if t.strip()]))
            translated_list = await self.translation_service.batch_translate(
                unique_texts, src_lang, tgt_lang, cache_only=True, translate_all=True
            )
            trans_map = dict(zip(unique_texts, translated_list))
            return self.pdf_reconstructor.reconstruct(pages_data, trans_map)

    async def _reconstruct_docx(self, content: bytes, src_lang: str, tgt_lang: str) -> bytes:
        """Translate DOCX run-by-run to preserve bold, italic, underline, fonts, and colours."""
        doc = DocxDocument(io.BytesIO(content))

        # Collect all runs from body paragraphs and table cells
        all_runs: list = []

        def _collect_runs(paragraphs):
            for para in paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        all_runs.append(run)

        for para in doc.paragraphs:
            _collect_runs([para])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _collect_runs(cell.paragraphs)

        if not all_runs:
            out = io.BytesIO()
            doc.save(out)
            return out.getvalue()

        # Translate all run texts concurrently
        run_texts = [run.text for run in all_runs]
        translated = await self.translation_service.batch_translate(
            run_texts, src_lang, tgt_lang, cache_only=True
        )

        # Write translated text back into each run.
        for run, trans_text in zip(all_runs, translated):
            if run.text.endswith(" ") and not trans_text.endswith(" "):
                trans_text += " "
            run.text = trans_text

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    async def _reconstruct_spreadsheet(self, content: bytes, src_lang: str, tgt_lang: str, file_type: str) -> bytes:
        if file_type == "csv":
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_excel(io.BytesIO(content))
            
        texts_to_translate = list(df.columns)
        for _, row in df.iterrows():
            for val in row.values:
                texts_to_translate.append(str(val) if pd.notna(val) else "")
                
        all_sentences = []
        cell_sentences = []
        for cell_text in texts_to_translate:
            sentences = self._split_into_sentences(cell_text)
            cell_sentences.append(sentences)
            all_sentences.extend(sentences)
        # CRITICAL FIX: Use the Knowledge Graph ONLY. Never hit the TMT API during reconstruction.
        # This prevents the 'Double API Hit' when downloading/viewing.
        translated = await self.translation_service.batch_translate(
            all_sentences, src_lang, tgt_lang, translate_all=True, cache_only=True
        )
        
        idx = 0
        translated_cells = []
        for sentences in cell_sentences:
            translated_cell_sentences = translated[idx:idx + len(sentences)]
            idx += len(sentences)
            translated_cells.append(" ".join(translated_cell_sentences))
        
        new_columns = translated_cells[:len(df.columns)]
        translated_vals = translated_cells[len(df.columns):]
        
        df.columns = new_columns
        idx = 0
        for i in range(len(df)):
            for j in range(len(df.columns)):
                if texts_to_translate[len(df.columns) + idx]:
                    if df.dtypes.iloc[j] != 'object':
                        df[df.columns[j]] = df[df.columns[j]].astype('object')
                    df.iat[i, j] = translated_vals[idx]
                idx += 1
                
        out = io.BytesIO()
        if file_type == "csv":
            df.to_csv(out, index=False)
        else:
            df.to_excel(out, index=False)
        return out.getvalue()

def get_document_processor(translation_service: TranslationService = Depends(get_translation_service)) -> DocumentProcessor:
    return DocumentProcessor(translation_service)
